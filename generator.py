Task:
Role: Automation / Business Systems
Responsibilities
Proposal draft generator (DOCX / structured output)
Template matching from past proposals
CRM auto-injection (GoHighLevel / Teltrans)
GoHighLevel:A CRM + marketing automation platform Used to manage leads, contacts, pipelines, emails, SMS, and workflows Common in agencies and SaaS businesses
Teltans: A CRM / workflow system used to track clients, proposals, and operational data. Similar idea: store records, update statuses, generate reports
Lead status updates & notes
Reporting outputs for staff review
Deliverables
Auto-generated proposal drafts
CRM entries created automatically
Basic reporting (what to bid on, what to skip)
Skills Used
Python
Document automation
Webhooks / APIs
Wehbooks: A webhook lets another system notify your code automatically. The CRM pushes data to you when something happens Example: lead status changes → your system gets notified

Think: “The system tells me when something happens.”
APIs stand for: Application Programming Interfaces) .An API lets your code talk to another system.Your program sends data (like a proposal or lead).The other system receives it and responds
Example: send a proposal → CRM creates a lead
Think: “I ask the system to do something.”
CRM workflows
code:
import requests
from docx import Document
import requests
# defines function called generate_proposal
#it allows you to generate many proposals with different titles/agencies. This function allows dynamic input instead hardcoding values
# so this shows how the function gives you flexbility of inputs
def generate_proposal(title, agency):
    return { # sends the values back whenever you call this function
        # here we use dictionary because Using a dictionary when you want data to be structured and labeled (formatted with meaning).
        "title": title,
        "agency": agency,
        "status": "Draft",
        # for the code below i weas building a proposal, which naturally has labeled sections:
        "sections": {
            "Executive Summary": (
                f"We propose to support {agency} by delivering services "
                f"aligned with the solicitation titled '{title}'."
            ),
            "Scope of Work": (
                "Provide planning, execution, and reporting services "
                "in accordance with government requirements."
            ),
            "Technical Approach": (
                "Use a structured, compliant, and scalable approach "
                "to meet all contract objectives."
            ),
            "Compliance Statement": (
                "All applicable federal requirements and standards will be met."
            )
        }
    }
# for here i created another function that has proposal and filename
# You need data first, then formatting.
def export_proposal_to_docx(proposal, filename):
    doc = Document() # starting blank document
    doc.add_heading(proposal["title"], level=1) # added heading with the title value
    # add paragrpah inside that document again i don't understand what is proposal agency
    doc.add_paragraph(f"Agency: {proposal['agency']}")
    doc.add_paragraph(f"Status: {proposal['status']}")# Same here

    for section, content in proposal["sections"].items():# so here it means by section and content is that document has sections and content what is ["sections"]
        doc.add_heading(section, level=2) # what happens in here
        doc.add_paragraph(content)

    doc.save(filename)

def select_template(agency):
    templates = {
        "Department of Defense": "dod_template.docx",
        "Department of Education": "education_template.docx"
    }
    return templates.get(agency, "default_template.docx")

def push_to_crm(proposal):
    payload = {
        "title": proposal["title"],
        "agency": proposal["agency"],
        "status": proposal["status"]
    }
    print("CRM payload ready:", payload)

def evaluate_bid(proposal):
    keywords = ["Defense", "Training", "Media"]
    score = sum(
        1 for word in keywords
        if word.lower() in proposal["title"].lower()
    )
    return "BID" if score >= 2 else "SKIP"

if __name__ == "__main__":
    proposal = generate_proposal(
        "Sports Media Training Services",
        "Department of Defense"
    )

    export_proposal_to_docx(proposal, "sports_media_proposal.docx")

    decision = evaluate_bid(proposal)
    push_to_crm(proposal)

    print("Bid decision:", decision)
    print("Proposal generated successfully.")






